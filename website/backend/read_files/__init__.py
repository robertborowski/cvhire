# ------------------------ imports start ------------------------
import PyPDF2
from docx import Document
# ------------------------ imports end ------------------------

# ------------------------ individual function start ------------------------
def get_file_contents_function(input_file, input_file_type):
  file_content = ''
  try:
    # ------------------------ format start ------------------------
    if input_file_type == '.pdf':
      pdf_reader = PyPDF2.PdfReader(input_file)
      for page in pdf_reader.pages:
        file_content += page.extract_text()
      return file_content
    # ------------------------ format end ------------------------
    # ------------------------ format start ------------------------
    elif input_file_type == '.docx':
      doc = Document(input_file)
      full_text = []
      # Extract content from the main body of the document
      for para in doc.paragraphs:
        full_text.append(para.text)
      # Extract content from the header and footer of the document
      for section in doc.sections:
        for para in section.first_page_header.paragraphs:
          full_text.append(para.text)
        for para in section.header.paragraphs:
          full_text.append(para.text)
        for para in section.first_page_footer.paragraphs:
          full_text.append(para.text)
        for para in section.footer.paragraphs:
          full_text.append(para.text)
      return '\n'.join(full_text)
    # ------------------------ format end ------------------------
    # ------------------------ format start ------------------------
    elif input_file_type == '.txt':
      file_content = input_file.read()
      return file_content
    # ------------------------ format end ------------------------
  except Exception as e:
    print(f'Error get_file_contents_function: {e}')
    pass
  return file_content
# ------------------------ individual function end ------------------------
